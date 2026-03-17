"""
VoteUnity — Patch existing report: replace placeholder boxes with real screenshots.
Then regenerate a fresh copy that includes images inline.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS = r"c:\Users\Lavya\Desktop\Online Voting System\screenshots"
OUT         = r"c:\Users\Lavya\Desktop\Online Voting System\VoteUnity_Report_Final.docx"

def img(name):
    p = os.path.join(SCREENSHOTS, name)
    return p if os.path.exists(p) else None

# Map every report section title -> (screenshot_file, caption)
SCREENSHOT_MAP = [
    # (unique text found in placeholder caption, screenshot filename, figure caption)
    ("Homepage",                    "01_homepage.png",               "Figure 5.1 — VoteUnity Homepage"),
    ("Voter Registration Page",     "02_register.png",               "Figure 5.2 — Voter Registration Form"),
    ("Voter Login Page",            "03_voter_login.png",            "Figure 5.3 — Voter Login with Face Verification"),
    ("Public Vote Verification",    "04_public_verify.png",          "Figure 5.7 — Public Vote Verification Page"),
    ("Admin Login — Step 1",        "05_admin_login_step1.png",      "Figure 5.8a — Admin Login (Step 1: Credentials)"),
    ("Admin Login — Step 2",        "06_admin_login_step2_face.png", "Figure 5.8b — Admin Login (Step 2: Face Verification)"),
    ("Admin Dashboard",             "07_admin_dashboard.png",        "Figure 5.9 — Admin Dashboard"),
    ("Admin Vote Audit",            "08_admin_view_votes.png",       "Figure 5.10 — Admin Vote Audit Trail"),
    ("Admin Audit Logs",            "09_admin_audit_logs.png",       "Figure 5.11 — Admin Audit Logs"),
    ("Admin — Manage Admins",       "10_admin_manage_admins.png",    "Figure 5.12 — Admin Manage Admins Panel"),
    ("Admin — Location Tracker",    "11_admin_location_tracker.png", "Figure 5.13 — Admin Location Tracker"),
    ("Tamper Demo",                 "12_admin_tamper_demo.png",      "Figure 5.14 — Tamper Detection Demo"),
    ("Registration",                "15_register_face_badge.png",    "Figure 5.2b — Face Recognition Engine Badge"),
    ("Face Recognition",            "16_homepage_features_section.png", "Figure 5.1b — Face Recognition Feature Card with Engine Badge"),
    ("Admin Register",              "17_admin_register.png",         "Figure 5.15 — Admin Registration Page"),
    ("Vote Casting",                "14_vote_redirect.png",          "Figure 5.4 — Vote Page (Requires Authentication)"),
]

# ── Re-run the full generator with images ────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

def set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    for attr in ['w:ascii','w:hAnsi','w:cs','w:eastAsia']:
        rFonts.set(qn(attr), 'Times New Roman')
    rpr.insert(0, rFonts)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text); set_font(r, 16, bold=True)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text); set_font(r, 14, bold=True)

def body(text, justify=True, bold=False, italic=False, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); set_font(r, 12, bold=bold, italic=italic)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); set_font(r, 12)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F2F2F2')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(9)
    rpr = r._r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    for a in ['w:ascii','w:hAnsi','w:cs']: rf.set(qn(a),'Courier New')
    rpr.insert(0, rf)

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'AAAAAA')
    pBdr.append(bot); pPr.append(pBdr)

def add_image(filename, caption, width=Inches(5.8)):
    path = img(filename)
    if path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run()
            run.add_picture(path, width=width)
        except Exception as e:
            print(f"  ! Could not embed {filename}: {e}")
            _placeholder(caption)
    else:
        _placeholder(caption)
    # Caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after  = Pt(14)
    cr = cp.add_run(caption); set_font(cr, 11, italic=True)

def _placeholder(caption):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0,0)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'D9D9D9')
    tcPr.append(shd)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30); p.paragraph_format.space_after = Pt(30)
    r = p.add_run(f'[ {caption} ]'); set_font(r, 11, italic=True, color=(80,80,80))

def tbl_hdr(tbl, headers, fill='C0C0C0'):
    row = tbl.rows[0]
    for i, h in enumerate(headers):
        p = row.cells[i].paragraphs[0]
        r = p.add_run(h); set_font(r, 11, bold=True)
        tc = row.cells[i]._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
        tcPr.append(shd)

def tbl_row(tbl, cells, bold=False):
    row = tbl.add_row()
    for i, text in enumerate(cells):
        p = row.cells[i].paragraphs[0]
        r = p.add_run(str(text)); set_font(r, 11, bold=bold)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()

for text, sz, bold, color in [
    ("PROJECT REPORT",              18, True,  None),
    ("VoteUnity",                   26, True,  (30,30,120)),
    ("Secure Online Voting System\nBiometric Auth · Blockchain Integrity · Admin Oversight",
                                    14, False, None),
]:
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(text); set_font(r, sz, bold=bold, color=color)

doc.add_paragraph()
for line in [
    "Technology Stack: PHP 8+  |  MySQL / PostgreSQL  |  Python  |  JavaScript",
    "Deployment: Vercel (Cloud)  |  XAMPP (Local)",
    "Academic Year: 2025 – 2026",
]:
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(line); set_font(lr, 12)

for _ in range(3): doc.add_paragraph()
for label, url in [("GitHub:", "https://github.com/lavyadamania/VoteUnity"),
                   ("Live App:", "https://vote-unity.vercel.app")]:
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = lp.add_run(f"{label}  "); set_font(r1, 12, bold=True)
    r2 = lp.add_run(url); set_font(r2, 12, color=(0,70,180))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Abstract")
body("VoteUnity is a full-stack, cloud-deployable online voting system that addresses security, "
     "transparency, and voter authentication in digital elections. It integrates three layers of "
     "biometric face recognition — ArcFace AI (cloud ONNX on Vercel), local ArcFace via Python "
     "DeepFace, and a PHP GD pixel-analysis fallback — ensuring robust voter identity verification "
     "at every sensitive step: registration, login, and vote casting.")
body("Vote integrity is guaranteed through a blockchain-inspired SHA-256 hash chain and Merkle tree. "
     "All vote data is encrypted at rest using AES-256-GCM. A comprehensive admin panel provides "
     "real-time statistics, a blockchain-style audit trail, GPS-based admin location tracking, "
     "role-based access with super-admin approval, and a tamper-detection demonstration. "
     "The system is deployed on Vercel with a Neon PostgreSQL backend, and supports local "
     "deployment via XAMPP with MySQL.")
hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. Introduction")
body("Elections are the cornerstone of democratic governance. Traditional paper-based voting is "
     "susceptible to fraud, logistical challenges, and low voter turnout due to geographic "
     "constraints. Online voting systems offer a promising solution only if they can guarantee "
     "authenticity, confidentiality, and tamper-proof record-keeping.")
body("VoteUnity demonstrates that a secure, end-to-end verifiable online voting system can be "
     "built using widely available technologies. By combining biometric authentication, "
     "cryptographic vote chaining, symmetric encryption, and a transparent public verification "
     "mechanism, VoteUnity provides a reference implementation suitable for academic study and "
     "demonstration.")
h2("2.1 Objectives")
for obj in [
    "Multi-factor voter authentication: Aadhaar-based identity, password, and biometric face verification.",
    "One-person-one-vote enforcement through database-level uniqueness constraints and session flags.",
    "Vote integrity using a blockchain-style SHA-256 hash chain and Merkle tree.",
    "Vote confidentiality via AES-256-GCM encryption at rest.",
    "Transparent, publicly accessible vote verification using receipt tokens and Merkle proofs.",
    "Admin panel with real-time statistics, audit logs, GPS tracking, and tamper detection.",
    "Support for both local (XAMPP/MySQL) and cloud (Vercel/PostgreSQL) deployment.",
    "Clear UI indicators showing which face recognition engine is actively processing biometric data.",
]: bullet(obj)

h2("2.2 Scope")
body("The project covers voter registration and authentication, candidate management, encrypted "
     "vote casting, a cryptographic audit trail, a public vote verification portal, and a "
     "role-based administration panel. It is scoped as a mock national-level election simulation "
     "using Indian political parties and Aadhaar numbers as identity proxies.")
hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. BACKGROUND
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Background")
h2("3.1 Challenges in Online Voting")
body("Online voting must simultaneously satisfy conflicting requirements: votes must be secret "
     "yet verifiable, and universal yet authenticated. Past systems suffered from weak "
     "authentication, lack of audit trails, centralised vulnerabilities, and inadequate "
     "transparency.")
h2("3.2 ArcFace Biometric Authentication")
body("ArcFace (Additive Angular Margin Loss, Deng et al. 2019) is the state-of-the-art face "
     "recognition method that trains neural networks to produce highly discriminative 512-dimensional "
     "face embeddings. The InsightFace MobileFaceNet model (w600k_mbf.onnx) achieves near-human "
     "accuracy on LFW benchmarks while remaining lightweight for serverless deployment. VoteUnity "
     "also supports DeepFace locally and falls back to a PHP GD pixel-difference comparison "
     "when Python is unavailable.")
h2("3.3 Blockchain-Style Hash Chaining")
body("Each vote stores the SHA-256 hash of its own content and the hash of the previous vote. "
     "Any alteration breaks the chain from that point forward, making tampering instantly "
     "detectable without needing a distributed ledger.")
h2("3.4 Merkle Trees for Vote Verification")
body("A Merkle tree allows any third party to verify that a specific vote is included in the "
     "vote set using only a logarithmic-size proof path rather than the entire dataset. "
     "This powers the public verification portal: voters can confirm their receipt without "
     "revealing other voters' data.")
h2("3.5 AES-256-GCM Vote Encryption")
body("Votes are encrypted using AES-256-GCM, providing both confidentiality and authenticated "
     "integrity. Each encryption uses a random 96-bit IV. The stored format is "
     "base64(IV):base64(tag):base64(ciphertext). The key is a 256-bit secret in an environment "
     "variable.")
h2("3.6 JWT Authentication")
body("Custom pure-PHP HMAC-SHA256 JWTs are issued alongside PHP sessions. Tokens contain "
     "user identity, role, and expiry, stored as HttpOnly Secure SameSite=Lax cookies.")
hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. SYSTEM DESCRIPTION
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. System Description")

h2("4.1 Technology Stack")
tbl = doc.add_table(rows=1, cols=3); tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_hdr(tbl, ['Layer', 'Technology', 'Details'])
for row in [
    ('Backend',             'PHP 8.0+',              'Server-rendered pages and business logic'),
    ('Database',            'MySQL / PostgreSQL',     'XAMPP local / Neon cloud via PDO'),
    ('Frontend',            'HTML5, CSS3, JS ES6',    'Dark-themed responsive UI'),
    ('Face Recog. (Cloud)', 'Python + ONNX Runtime',  'ArcFace MobileFaceNet on Vercel serverless'),
    ('Face Recog. (Local)', 'Python + DeepFace',      'ArcFace via Python subprocess'),
    ('Face Recog. (FB)',    'PHP GD Library',          'Pixel-diff grayscale fallback 64x64'),
    ('Maps',                'Leaflet.js 1.9.4',         'Admin GPS location visualisation'),
    ('Auth',                'PHP Sessions + JWT',       'HMAC-SHA256, HttpOnly cookies'),
    ('Encryption',          'AES-256-GCM (OpenSSL)',    'Vote data encrypted at rest'),
    ('Integrity',           'SHA-256 + Merkle Tree',    'Hash chain + inclusion proofs'),
    ('Deployment',          'Vercel + Heroku',          'PHP + Python serverless runtimes'),
]: tbl_row(tbl, row)
doc.add_paragraph()

h2("4.2 System Architecture — Block Diagram")
body("The following block diagram shows the high-level architecture of VoteUnity:", after=4)

arch_tbl = doc.add_table(rows=1, cols=1)
arch_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = arch_tbl.cell(0,0)
tc = cell._tc; tcPr = tc.get_or_add_tcPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F0F0F0')
tcPr.append(shd)
arch_lines = [
    "  +------------------------------------------------------------------+",
    "  |                     USER / BROWSER                              |",
    "  |  Voter UI  <---------------------------------------------> Admin UI  |",
    "  +-----------------------------+------------------------------------+",
    "                                |  HTTPS",
    "  +-----------------------------v------------------------------------+",
    "  |                  PHP APPLICATION LAYER                          |",
    "  |  includes/          |  pages/voter/       |  pages/admin/       |",
    "  |  - functions.php    |  - register.php     |  - dashboard.php    |",
    "  |  - jwt_helper.php   |  - login.php        |  - view_votes.php   |",
    "  |  - encryption.php   |  - vote.php         |  - audit_logs.php   |",
    "  |  - merkle_tree.php  |  - verify.php       |  - manage_admins    |",
    "  |  - audit_logger.php |                     |  - location_tracker |",
    "  +--------+-------------------------------+-----------+-------------+",
    "           |                               |           |",
    "  +--------v--------+    +----------------v---+  +----v-----------+",
    "  |   DATABASE      |    |  FACE RECOGNITION  |  |  AUDIT / JWT   |",
    "  |  MySQL /        |    |  1. ArcFace API    |  |  audit_logs    |",
    "  |  PostgreSQL     |    |     (Vercel ONNX)  |  |  JWT tokens    |",
    "  |  - users        |    |  2. ArcFace Local  |  |  Sessions      |",
    "  |  - votes        |    |     (DeepFace)     |  +----------------+",
    "  |  - candidates   |    |  3. PHP GD         |",
    "  |  - admins       |    |     (Fallback)     |",
    "  |  - audit_logs   |    +--------------------+",
    "  |  - admin_locs   |",
    "  +-----------------+",
]
for line in arch_lines:
    ap = cell.add_paragraph()
    ar = ap.add_run(line)
    ar.font.name = "Courier New"; ar.font.size = Pt(8)
    ap.paragraph_format.space_before = Pt(0); ap.paragraph_format.space_after = Pt(0)
cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(4)
cr = cap.add_run("Figure 1: VoteUnity System Architecture Block Diagram"); set_font(cr, 11, italic=True)
doc.add_paragraph()

h2("4.3 Database Schema")
body("VoteUnity uses six database tables:", after=4)
db_tbl = doc.add_table(rows=1, cols=3); db_tbl.style = 'Table Grid'
db_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_hdr(db_tbl, ['Table', 'Key Columns', 'Purpose'])
for row in [
    ('users','id, name, email, password, aadhaar_number, face_image, has_voted','Registered voter accounts'),
    ('candidates','id, name, party, symbol, photo','Election candidates'),
    ('votes','id, user_id, candidate_id, vote_hash, previous_hash, encrypted_vote, block_index, nonce, merkle_root, vote_receipt','Immutable vote ledger'),
    ('admins','id, username, password, face_image, is_super_admin, is_approved, can_view_votes, can_reset_votes, ...','Admin accounts with permissions'),
    ('audit_logs','id, event_type, actor_type, actor_id, details, ip_address, user_agent, created_at','Security event trail'),
    ('admin_locations','id, admin_id, latitude, longitude, accuracy, ip_address, tracked_at','GPS admin login tracking'),
]: tbl_row(db_tbl, row)
doc.add_paragraph()

h2("4.4 Key Modules")
for title, source, desc in [
    ("4.4.1  Voter Registration", "pages/register.php",
     "Collects name, email, 12-digit Aadhaar, password, and webcam face photo. The face recognition "
     "engine badge is shown before capture, indicating which algorithm will process biometric data. "
     "Face stored as base64 in PostgreSQL (Vercel) or as JPEG file in uploads/faces/ (local)."),
    ("4.4.2  Voter Authentication", "pages/login.php",
     "Two-factor: email + bcrypt password, then live face capture compared via compareFaces(). "
     "A 60% similarity threshold must be met. On success, PHP session and signed JWT are established. "
     "Engine name is shown in the UI and included in all error / success messages."),
    ("4.4.3  Vote Casting", "pages/vote.php",
     "Three phases: (1) mandatory second face verification, (2) candidate selection from responsive "
     "grid, (3) atomic DB transaction generating SHA-256 vote hash, AES-256-GCM encrypted vote data, "
     "Merkle root recomputation, and unique receipt token issuance."),
    ("4.4.4  Public Vote Verification", "pages/verify.php",
     "No-login public page showing aggregate statistics and vote distribution. "
     "Voters enter their receipt token to receive a Merkle inclusion proof without revealing "
     "candidate choice or voter identity."),
    ("4.4.5  Admin Authentication", "pages/admin/login.php",
     "Two-step: credentials then biometric face verification. Non-super admins also provide "
     "GPS coordinates via browser Geolocation API before session activation."),
    ("4.4.6  Admin Dashboard", "pages/admin/dashboard.php",
     "Real-time stats: voter count, votes cast, turnout, candidate count, per-candidate chart, "
     "recent votes, chain integrity status, Merkle root, and active face recognition engine badge."),
    ("4.4.7  Vote Audit Trail", "pages/admin/view_votes.php",
     "Blockchain-style chain visualization. Each card shows block index, vote hash, previous hash, "
     "and chain validity. A genesis block heads the chain. Broken links show in red."),
    ("4.4.8  Audit Logs", "pages/admin/audit_logs.php + includes/audit_logger.php",
     "All security events logged: LOGIN, LOGIN_FAIL, LOGOUT, VOTE_CAST, VOTE_TAMPER, "
     "ADMIN_ACTION, CHAIN_VERIFY, REGISTER, FACE_VERIFY, SYSTEM. Filterable, paginated viewer."),
    ("4.4.9  Admin Management", "pages/admin/manage_admins.php",
     "Super admin approves/rejects pending admins, sets granular permissions "
     "(view votes, manage candidates, reset votes, manage admins), and revokes access."),
    ("4.4.10  Location Tracking", "pages/admin/location_tracker.php + api_location.php",
     "Leaflet.js map with markers at all admin login GPS coordinates. Each marker shows "
     "admin username, timestamp, IP address, and accuracy radius."),
    ("4.4.11  Tamper Demonstration", "pages/admin/tamper_demo.php",
     "Presentation tool: simulates vote tampering, then immediately shows broken hash chain "
     "blocks in red. Logs a VOTE_TAMPER audit event. Includes vote reset for demo resets."),
    ("4.4.12  Face Engine Detection", "includes/functions.php — detectFaceRecognitionMethod()",
     "Detects active engine at runtime: arcface_api (Vercel), arcface_local (Python present), "
     "gd_strict (PHP GD), or fallback. getFaceMethodInfo() returns label, description, colour, "
     "icon, and tier for display in UI on all face-capture pages."),
]:
    h2(title)
    fp = doc.add_paragraph()
    fp.paragraph_format.space_after = Pt(4)
    r1 = fp.add_run("Source: "); set_font(r1, 12, bold=True)
    r2 = fp.add_run(source); set_font(r2, 12, italic=True, color=(60,60,140))
    body(desc)

hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. RESULTS AND DISCUSSIONS
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Results and Discussions")
body("VoteUnity was deployed on Vercel with a Neon PostgreSQL backend. All core workflows — "
     "voter registration, login with face verification, vote casting, receipt-based verification, "
     "and admin oversight — function correctly. Screenshots below are taken from the live "
     "deployment at https://vote-unity.vercel.app.")

h2("5.1  Homepage")
body("The homepage presents VoteUnity's features including the active face recognition engine badge "
     "(e.g., 'ArcFace AI (Cloud) — Premium' in green) on the Face Recognition feature card. "
     "CTA buttons adapt based on login state.")
add_image("01_homepage.png", "Figure 5.1a — VoteUnity Homepage (full page)")
add_image("16_homepage_features_section.png", "Figure 5.1b — Feature cards with active Face Recognition Engine badge", width=Inches(5.5))

h2("5.2  Voter Registration")
body("The registration form collects name, email, Aadhaar, password, and webcam face photo. "
     "A face recognition engine badge above the webcam clearly indicates which algorithm will "
     "process biometric data.")
add_image("02_register.png", "Figure 5.2a — Voter Registration Form")
add_image("15_register_face_badge.png", "Figure 5.2b — Face Recognition Engine Badge (scrolled view)", width=Inches(5.5))

h2("5.3  Voter Login with Face Verification")
body("After email/password validation, the voter captures their face via webcam. The face method "
     "box shows the active engine and description. Flash messages include similarity score and "
     "engine name on both success and failure.")
add_image("03_voter_login.png", "Figure 5.3 — Voter Login Page with Face Verification Section")

h2("5.4  Vote Casting — Authentication Gate")
body("The vote page enforces authentication. Unauthenticated users are redirected to login with "
     "a security message. After login, users must pass a second independent face verification "
     "before the candidate grid is shown.")
add_image("14_vote_redirect.png", "Figure 5.4 — Vote Page Access (Security Redirect for Unauthenticated Users)")

h2("5.5  Public Vote Verification")
body("No-login public page showing total voters, votes cast, turnout, Merkle root, and chain "
     "integrity status. A receipt input verifies individual votes with Merkle inclusion proofs "
     "without revealing voter identity or candidate choice.")
add_image("04_public_verify.png", "Figure 5.5 — Public Vote Verification Page")

h2("5.6  Admin Login — Step 1: Credentials")
body("Admin login starts with username and password verification against the database. "
     "Only approved admins can proceed.")
add_image("05_admin_login_step1.png", "Figure 5.6 — Admin Login Step 1 (Credentials)")

h2("5.7  Admin Login — Step 2: Face Verification")
body("After credentials, the admin must pass face verification. The active face recognition "
     "engine badge is displayed. For first-time logins, the face is stored. Subsequent logins "
     "require a 60% match against the stored face. Non-super admins then share GPS coordinates.")
add_image("06_admin_login_step2_face.png", "Figure 5.7 — Admin Login Step 2 (Face Verification with Engine Badge)")

h2("5.8  Admin Registration")
body("New admins register with username and password. Their account is created with "
     "is_approved = false and must be approved by a super admin before they can log in.")
add_image("17_admin_register.png", "Figure 5.8 — Admin Registration Page")

h2("5.9  Admin Dashboard")
body("The admin dashboard displays Registered Voters, Votes Cast, Turnout %, Candidate count, "
     "per-candidate vote distribution chart, 10 most recent votes, chain integrity status, "
     "Merkle root, and the active face recognition engine badge. Protected — requires admin session.")
add_image("07_admin_dashboard.png", "Figure 5.9 — Admin Dashboard (redirects to login when unauthenticated)")

h2("5.10  Vote Audit Trail")
body("Blockchain-style visualization of all votes. Each block shows index, hash, previous hash, "
     "and chain validity indicator. Tampered blocks are highlighted in red.")
add_image("08_admin_view_votes.png", "Figure 5.10 — Admin Vote Audit Trail")

h2("5.11  Audit Logs")
body("Filterable, paginated log of all security events with colour-coded event type badges, "
     "actor details, IP address, and timestamp.")
add_image("09_admin_audit_logs.png", "Figure 5.11 — Admin Audit Logs")

h2("5.12  Manage Admins")
body("Super admin panel to approve/reject pending admins and set granular permissions per admin.")
add_image("10_admin_manage_admins.png", "Figure 5.12 — Manage Admins Panel")

h2("5.13  Location Tracker")
body("Leaflet.js map with markers at all admin login GPS coordinates, with username, timestamp, "
     "IP, and accuracy for each marker.")
add_image("11_admin_location_tracker.png", "Figure 5.13 — Admin Location Tracker Map")

h2("5.14  Tamper Detection Demo")
body("Simulates vote tampering: directly modifies a vote in the database then immediately shows "
     "the broken chain blocks. A VOTE_TAMPER event is logged. Reset button clears all votes.")
add_image("12_admin_tamper_demo.png", "Figure 5.14 — Tamper Detection Demonstration")

h2("5.15  Face Recognition Engine Performance")
body("Three face recognition methods are integrated with the following characteristics:")
perf_tbl = doc.add_table(rows=1, cols=4); perf_tbl.style = 'Table Grid'
perf_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_hdr(perf_tbl, ['Method','Threshold','Accuracy','Deployment'])
for row in [
    ('ArcFace AI (Cloud)', '0.40 cosine', 'High — deep learning 512-dim embeddings', 'Vercel serverless'),
    ('ArcFace AI (Local)', '60% similarity','High — DeepFace + ArcFace model',        'XAMPP + Python 3'),
    ('GD Pixel Analysis',  '60% similarity','Basic — pixel-diff, lighting sensitive', 'Any PHP + GD ext.'),
    ('Bypass (Fallback)',   'None',          'None — always passes',                  'Emergency only'),
]: tbl_row(perf_tbl, row)
doc.add_paragraph()
body("The active engine is auto-detected at runtime and shown as a colour-coded badge: "
     "green for ArcFace Premium, amber for GD Basic, red for Bypass — on all pages where "
     "face capture occurs.")
hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. CONCLUSION AND FUTURE SCOPE
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Conclusion and Future Scope")
h2("6.1 Conclusion")
body("VoteUnity successfully demonstrates a secure, full-stack online voting system that addresses "
     "the fundamental requirements of modern e-voting: strong voter authentication (biometric + "
     "Aadhaar + password), vote confidentiality (AES-256-GCM), vote integrity (SHA-256 hash "
     "chain + Merkle tree), public verifiability (receipt tokens), and administrative accountability "
     "(audit logs, GPS tracking, role-based access).")
body("The three-tier fallback face recognition architecture ensures the system remains functional "
     "across different deployment environments while always informing users of the active engine "
     "through prominent, colour-coded badges. The blockchain-style audit trail and tamper "
     "demonstration prove that any vote record modification is immediately detectable.")
h2("6.2 Future Scope")
for item in [
    "Zero-Knowledge Proofs (ZKPs): Mathematically stronger privacy guarantees for vote receipt verification.",
    "End-to-End Verifiable (E2E) Voting: Implement Helios or ElectionGuard for independent cryptographic verification.",
    "Multi-Election Support: Multiple simultaneous elections with different voter pools and periods.",
    "Liveness Detection: Anti-spoofing (blink detection, head movement) to prevent photo-spoofing attacks.",
    "Real Aadhaar API Integration: Replace simulated validation with UIDAI API for production deployments.",
    "Mobile Application: React Native or Flutter app with hardware-accelerated face recognition.",
    "Distributed Ledger: Migrate hash chain to Hyperledger Fabric to eliminate single-point compromise.",
    "Homomorphic Encryption: Vote tallying without decrypting individual ballots for stronger ballot secrecy.",
    "TOTP MFA: Add Time-based One-Time Password as additional admin authentication factor.",
    "Accessibility Improvements: Screen reader support, keyboard navigation, alternative identity verification.",
]: bullet(item)
hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. References")
for i, ref in enumerate([
    "Deng, J., Guo, J., Xue, N., & Zafeiriou, S. (2019). ArcFace: Additive Angular Margin Loss for Deep Face Recognition. CVPR 2019, pp. 4690–4699.",
    "Nakamoto, S. (2008). Bitcoin: A Peer-to-Peer Electronic Cash System. https://bitcoin.org/bitcoin.pdf",
    "Merkle, R. C. (1987). A Digital Signature Based on a Conventional Encryption Function. CRYPTO '87, LNCS vol. 293, Springer.",
    "Jones, D. W. (2003). The Case for Cryptographic Voter Verification. USENIX Workshop EVT 2003.",
    "NIST. (2001). FIPS PUB 197: Advanced Encryption Standard (AES).",
    "NIST. (2007). FIPS PUB 180-4: Secure Hash Standard (SHS).",
    "IETF. (2015). RFC 7519: JSON Web Token (JWT). https://datatracker.ietf.org/doc/html/rfc7519",
    "Serengil, S. I. & Ozpinar, A. (2020). LightFace: A Hybrid Deep Face Recognition Framework. ASYU 2020, IEEE.",
    "OWASP Foundation. (2021). OWASP Top Ten. https://owasp.org/www-project-top-ten/",
    "InsightFace. (2023). InsightFace: State-of-the-art 2D and 3D Face Analysis. https://github.com/deepinsight/insightface",
    "Vercel Inc. (2024). Vercel Documentation: Serverless Functions. https://vercel.com/docs/functions",
    "The PHP Group. (2023). PHP 8.2 Manual. https://www.php.net/manual/en/",
    "PostgreSQL Global Development Group. (2024). PostgreSQL 16 Documentation. https://www.postgresql.org/docs/",
    "Leaflet.js Contributors. (2023). Leaflet — An open-source JS library for maps. https://leafletjs.com/",
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    r = p.add_run(f"[{i}]  {ref}"); set_font(r, 11)
hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. CODE SNIPPETS
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Code — Key Implementation Snippets")

h2("8.1  Face Recognition Engine Detection")
body("Source: includes/functions.php", bold=True, after=4)
code("""function detectFaceRecognitionMethod() {
    $isVercel = getenv('VERCEL') || getenv('VERCEL_URL');
    if ($isVercel) return 'arcface_api';

    $pythonScript = dirname(__DIR__) . '/python/arcface_verify.py';
    if (file_exists($pythonScript)) {
        foreach (['python3', 'python'] as $candidate) {
            $check = @shell_exec($candidate . ' --version 2>&1');
            if ($check && stripos($check, 'python') !== false)
                return 'arcface_local';
        }
    }
    if (extension_loaded('gd')) return 'gd_strict';
    return 'fallback';
}

function getFaceMethodInfo($method = null) {
    if ($method === null) $method = detectFaceRecognitionMethod();
    $methods = [
        'arcface_api'   => ['label'=>'ArcFace AI (Cloud)', 'icon'=>'Brain',
                            'color'=>'#10b981', 'tier'=>'Premium',
                            'description'=>'Serverless ArcFace ONNX via Vercel Python'],
        'arcface_local' => ['label'=>'ArcFace AI (Local)', 'icon'=>'Brain',
                            'color'=>'#6366f1', 'tier'=>'Premium',
                            'description'=>'ArcFace via local Python + DeepFace'],
        'gd_strict'     => ['label'=>'GD Pixel Analysis', 'icon'=>'Grid',
                            'color'=>'#f59e0b', 'tier'=>'Basic',
                            'description'=>'PHP GD pixel-diff grayscale 64x64'],
        'fallback'      => ['label'=>'Bypass (No Engine)', 'icon'=>'Warning',
                            'color'=>'#ef4444', 'tier'=>'None',
                            'description'=>'No face recognition engine available'],
    ];
    return $methods[$method] ?? $methods['fallback'];
}""")

h2("8.2  Vote Casting — Blockchain Transaction")
body("Source: pages/vote.php", bold=True, after=4)
code("""$previousHash = getLastVoteHash($pdo);   // 'GENESIS' or last hash
$blockIndex   = $pdo->query("SELECT COUNT(*) FROM votes")->fetchColumn() + 1;
$nonce        = bin2hex(random_bytes(16));

$voteHash = hash('sha256',
    $_SESSION['user_id'] . $candidateId . $timestamp . $previousHash . $nonce
);

$encryptedVote = encryptVote(json_encode([
    'user_id'=>$_SESSION['user_id'], 'candidate_id'=>$candidateId,
    'timestamp'=>$timestamp, 'block_index'=>$blockIndex
]));   // AES-256-GCM

$voteReceipt = hash('sha256', random_bytes(32) . $voteHash . $timestamp);

// Inside transaction
$pdo->prepare("INSERT INTO votes (user_id, candidate_id, vote_hash, previous_hash,
    encrypted_vote, block_index, nonce, vote_receipt, timestamp) VALUES (?,?,?,?,?,?,?,?,?)")
    ->execute([...]);

$merkleRoot = computeMerkleRoot($pdo)['root'];
$pdo->prepare("UPDATE votes SET merkle_root=? WHERE vote_hash=?")
    ->execute([$merkleRoot, $voteHash]);
$pdo->prepare("UPDATE users SET has_voted=TRUE WHERE id=?")->execute([$_SESSION['user_id']]);
$pdo->commit();""")

h2("8.3  AES-256-GCM Encryption/Decryption")
body("Source: includes/encryption.php", bold=True, after=4)
code("""function encryptVote($plaintext) {
    $key    = hex2bin(getEncryptionKey());   // 32-byte key
    $iv     = random_bytes(12);              // 96-bit IV
    $tag    = '';
    $cipher = openssl_encrypt($plaintext, 'aes-256-gcm', $key,
                              OPENSSL_RAW_DATA, $iv, $tag);
    return base64_encode($iv) . ':' . base64_encode($tag) . ':' . base64_encode($cipher);
}

function decryptVote($encrypted) {
    [$iv, $tag, $cipher] = array_map('base64_decode', explode(':', $encrypted));
    $key = hex2bin(getEncryptionKey());
    return openssl_decrypt($cipher, 'aes-256-gcm', $key, OPENSSL_RAW_DATA, $iv, $tag);
}""")

h2("8.4  Merkle Tree — Build and Proof")
body("Source: includes/merkle_tree.php", bold=True, after=4)
code("""function computeMerkleRoot($pdo) {
    $hashes = $pdo->query("SELECT vote_hash FROM votes ORDER BY block_index")
                  ->fetchAll(PDO::FETCH_COLUMN);
    if (empty($hashes)) return ['root'=>str_repeat('0',64), 'leaves'=>[]];

    $level = $hashes;
    while (count($level) > 1) {
        if (count($level) % 2 !== 0) $level[] = end($level); // duplicate last if odd
        $next = [];
        for ($i = 0; $i < count($level); $i += 2)
            $next[] = hash('sha256', $level[$i] . $level[$i+1]);
        $level = $next;
    }
    return ['root'=>$level[0], 'leaves'=>$hashes];
}""")

h2("8.5  ArcFace Serverless (Python — Vercel)")
body("Source: api/verify_face.py", bold=True, after=4)
code("""import onnxruntime as ort
import numpy as np

_session = None

def load_model():
    global _session
    if _session is None:
        if not os.path.exists(MODEL_PATH):
            download_model()          # Downloads w600k_mbf.onnx to /tmp on cold start
        _session = ort.InferenceSession(MODEL_PATH)
    return _session

def get_embedding(img_array):
    session = load_model()
    img     = preprocess(img_array)       # 112x112, normalised to [-1, 1]
    output  = session.run(None, {session.get_inputs()[0].name: img})[0][0]
    return output / np.linalg.norm(output)   # L2-normalise

def handler(request):
    body  = json.loads(request.body)
    emb1  = get_embedding(decode_image(body['image1']))
    emb2  = get_embedding(decode_image(body['image2']))
    score = float(np.dot(emb1, emb2))       # cosine similarity (already L2-normalised)
    return {'match': score >= 0.40, 'score': round(score, 4), 'method': 'arcface_api'}""")
hline()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. PROJECT LINKS
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. Project Links")
link_tbl = doc.add_table(rows=1, cols=2); link_tbl.style = 'Table Grid'
link_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_hdr(link_tbl, ['Resource', 'URL'])
for res, url in [
    ('GitHub Repository',           'https://github.com/lavyadamania/VoteUnity'),
    ('Live Application (Vercel)',    'https://vote-unity.vercel.app'),
    ('Admin Panel',                  'https://vote-unity.vercel.app/pages/admin/login.php'),
    ('Public Vote Verification',     'https://vote-unity.vercel.app/pages/verify.php'),
    ('Admin Credentials',            'Username: admin  |  Password: admin123'),
]:
    row = link_tbl.add_row()
    p0 = row.cells[0].paragraphs[0]; r0 = p0.add_run(res); set_font(r0, 12, bold=True)
    p1 = row.cells[1].paragraphs[0]; r1 = p1.add_run(url)
    set_font(r1, 12, color=(0,70,180) if url.startswith('http') else None)
doc.add_paragraph()

ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = ep.add_run("— End of Report —"); set_font(er, 12, italic=True, color=(100,100,100))

doc.save(OUT)
print(f"Report saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)//1024} KB")
